"""
Document Processor for KERBERUS Tabular Review.

Handles parsing of PDF and DOCX files, extracting text with page references.
"""

import logging
import os
from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple
from dataclasses import dataclass, field
import hashlib

logger = logging.getLogger(__name__)


@dataclass
class DocumentPage:
    """Represents a single page of a document."""
    page_number: int
    text: str
    char_start: int  # Character offset in full document
    char_end: int


@dataclass
class ParsedDocument:
    """Represents a fully parsed document."""
    document_id: str
    filename: str
    file_type: str
    total_pages: int
    full_text: str
    pages: List[DocumentPage]
    file_hash: str
    file_size: int
    
    def get_page_for_position(self, char_position: int) -> Optional[int]:
        """Get page number for a character position."""
        for page in self.pages:
            if page.char_start <= char_position < page.char_end:
                return page.page_number
        return None
    
    def get_context_around(self, search_text: str, context_chars: int = 200) -> Optional[Dict]:
        """
        Find text in document and return context with page reference.
        
        Args:
            search_text: Text to find
            context_chars: Characters of context around match
            
        Returns:
            Dict with page, quote, and position info
        """
        # Normalize both texts for matching
        normalized_full = ' '.join(self.full_text.split())
        normalized_search = ' '.join(search_text.split())
        
        pos = normalized_full.lower().find(normalized_search.lower())
        if pos == -1:
            # Try fuzzy match with first 50 chars
            short_search = normalized_search[:50]
            pos = normalized_full.lower().find(short_search.lower())
            
        if pos == -1:
            return None
            
        # Get context
        start = max(0, pos - context_chars)
        end = min(len(normalized_full), pos + len(normalized_search) + context_chars)
        quote = normalized_full[start:end]
        
        # Add ellipsis if truncated
        if start > 0:
            quote = "..." + quote
        if end < len(normalized_full):
            quote = quote + "..."
            
        # Find page (approximate since we normalized)
        page_num = self.get_page_for_position(pos)
        
        return {
            "page": page_num or 1,
            "quote": quote,
            "position": pos
        }


class DocumentProcessor:
    """
    Process documents for tabular review extraction.
    
    Supports:
    - PDF files (via pymupdf/fitz)
    - DOCX files (via python-docx)
    - TXT files (direct read)
    """
    
    MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB limit
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.doc', '.txt'}
    
    def __init__(self):
        """Initialize document processor."""
        self._check_dependencies()
    
    def _check_dependencies(self):
        """Check that required libraries are available."""
        try:
            import fitz  # pymupdf
            self._has_pdf = True
        except ImportError:
            logger.warning("pymupdf not installed - PDF support disabled")
            self._has_pdf = False
            
        try:
            import docx
            self._has_docx = True
        except ImportError:
            logger.warning("python-docx not installed - DOCX support disabled")
            self._has_docx = False
    
    def parse_file(self, file_path: str) -> ParsedDocument:
        """
        Parse a document file and extract text with page references.
        
        Args:
            file_path: Path to the file
            
        Returns:
            ParsedDocument with full text and page-level breakdown
            
        Raises:
            ValueError: If file type not supported or file too large
            FileNotFoundError: If file doesn't exist
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
            
        # Check file size
        file_size = path.stat().st_size
        if file_size > self.MAX_FILE_SIZE:
            raise ValueError(f"File too large: {file_size / 1024 / 1024:.1f}MB (max {self.MAX_FILE_SIZE / 1024 / 1024}MB)")
            
        # Check extension
        ext = path.suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {ext}. Supported: {self.SUPPORTED_EXTENSIONS}")
            
        # Calculate file hash for deduplication
        file_hash = self._calculate_hash(path)
        
        # Parse based on type
        if ext == '.pdf':
            return self._parse_pdf(path, file_hash, file_size)
        elif ext in {'.docx', '.doc'}:
            return self._parse_docx(path, file_hash, file_size)
        elif ext == '.txt':
            return self._parse_txt(path, file_hash, file_size)
        else:
            raise ValueError(f"Unsupported file type: {ext}")
    
    def parse_bytes(self, content: bytes, filename: str) -> ParsedDocument:
        """
        Parse document from bytes (for uploaded files).
        
        Args:
            content: File content as bytes
            filename: Original filename
            
        Returns:
            ParsedDocument
        """
        import tempfile
        
        ext = Path(filename).suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {ext}")
            
        # Write to temp file and parse
        with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
            tmp.write(content)
            tmp_path = tmp.name
            
        try:
            doc = self.parse_file(tmp_path)
            doc.filename = filename  # Use original filename
            return doc
        finally:
            # Cleanup temp file
            os.unlink(tmp_path)
    
    def _calculate_hash(self, path: Path) -> str:
        """Calculate SHA256 hash of file."""
        sha256 = hashlib.sha256()
        with open(path, 'rb') as f:
            for chunk in iter(lambda: f.read(8192), b''):
                sha256.update(chunk)
        return sha256.hexdigest()[:16]  # First 16 chars for brevity
    
    def _parse_pdf(self, path: Path, file_hash: str, file_size: int) -> ParsedDocument:
        """Parse PDF file using pymupdf."""
        if not self._has_pdf:
            raise ImportError("pymupdf not installed. Run: pip install pymupdf")
            
        import fitz
        
        pages = []
        full_text_parts = []
        char_offset = 0
        
        doc = fitz.open(str(path))
        
        try:
            for page_num, page in enumerate(doc, start=1):
                text = page.get_text()
                
                pages.append(DocumentPage(
                    page_number=page_num,
                    text=text,
                    char_start=char_offset,
                    char_end=char_offset + len(text)
                ))
                
                full_text_parts.append(text)
                char_offset += len(text) + 1  # +1 for newline between pages
                
            full_text = "\n".join(full_text_parts)
            
            return ParsedDocument(
                document_id=file_hash,
                filename=path.name,
                file_type="pdf",
                total_pages=len(pages),
                full_text=full_text,
                pages=pages,
                file_hash=file_hash,
                file_size=file_size
            )
            
        finally:
            doc.close()
    
    def _parse_docx(self, path: Path, file_hash: str, file_size: int) -> ParsedDocument:
        """Parse DOCX file using python-docx."""
        if not self._has_docx:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
            
        from docx import Document
        
        doc = Document(str(path))
        
        # DOCX doesn't have pages in the same way, treat as single page
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
                
        full_text = "\n\n".join(paragraphs)
        
        # Create single page
        pages = [DocumentPage(
            page_number=1,
            text=full_text,
            char_start=0,
            char_end=len(full_text)
        )]
        
        # Estimate pages based on character count (approx 3000 chars per page)
        estimated_pages = max(1, len(full_text) // 3000)
        
        return ParsedDocument(
            document_id=file_hash,
            filename=path.name,
            file_type="docx",
            total_pages=estimated_pages,
            full_text=full_text,
            pages=pages,
            file_hash=file_hash,
            file_size=file_size
        )
    
    def _parse_txt(self, path: Path, file_hash: str, file_size: int) -> ParsedDocument:
        """Parse plain text file."""
        with open(path, 'r', encoding='utf-8', errors='replace') as f:
            full_text = f.read()
            
        pages = [DocumentPage(
            page_number=1,
            text=full_text,
            char_start=0,
            char_end=len(full_text)
        )]
        
        return ParsedDocument(
            document_id=file_hash,
            filename=path.name,
            file_type="txt",
            total_pages=1,
            full_text=full_text,
            pages=pages,
            file_hash=file_hash,
            file_size=file_size
        )
    
    def chunk_for_extraction(
        self,
        doc: ParsedDocument,
        max_chunk_size: int = 4000,
        overlap: int = 200
    ) -> List[Dict[str, Any]]:
        """
        Split document into chunks for LLM processing.
        
        Args:
            doc: Parsed document
            max_chunk_size: Maximum characters per chunk
            overlap: Character overlap between chunks
            
        Returns:
            List of chunks with text and metadata
        """
        chunks = []
        text = doc.full_text
        
        if len(text) <= max_chunk_size:
            # Small document, single chunk
            return [{
                "chunk_index": 0,
                "text": text,
                "char_start": 0,
                "char_end": len(text),
                "is_full_document": True
            }]
            
        # Split into overlapping chunks
        start = 0
        chunk_index = 0
        
        while start < len(text):
            end = start + max_chunk_size
            
            # Try to break at paragraph boundary
            if end < len(text):
                # Look for paragraph break near end
                para_break = text.rfind('\n\n', start + max_chunk_size - 500, end)
                if para_break > start:
                    end = para_break
                    
            chunk_text = text[start:end]
            
            chunks.append({
                "chunk_index": chunk_index,
                "text": chunk_text,
                "char_start": start,
                "char_end": end,
                "is_full_document": False
            })
            
            # Move start with overlap
            start = end - overlap
            chunk_index += 1
            
        return chunks
